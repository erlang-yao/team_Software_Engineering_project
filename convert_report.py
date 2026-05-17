# -*- coding: utf-8 -*-
"""将 Markdown 实验报告转为 Word 文档"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

with open('g:/team_Software_Engineering_project/实验报告-Web UI接入与前后端分离.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_lines = []
in_table = False
table_rows = []

def add_table_to_doc(rows):
    if not rows:
        return
    # 解析表头和数据行
    header_cells = [c.strip() for c in rows[0].split('|')[1:-1]]
    data = []
    for r in rows[2:]:  # 跳过表头分隔行
        cells = [c.strip() for c in r.split('|')[1:-1]]
        if cells:
            data.append(cells)
    if not data:
        return
    table = doc.add_table(rows=1 + len(data), cols=len(header_cells))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # 数据
    for ri, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.rows[ri + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

def flush_code():
    global code_lines
    if code_lines:
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.left_indent = Cm(1)
        code_lines = []

def flush_table():
    global table_rows
    if table_rows:
        add_table_to_doc(table_rows)
        table_rows = []

while i < len(lines):
    line = lines[i]

    # 代码块
    if line.startswith('```'):
        flush_table()
        if in_code_block:
            flush_code()
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    if in_code_block:
        code_lines.append(line.rstrip('\n'))
        i += 1
        continue

    # 表格行
    if line.startswith('|') and line.strip().endswith('|'):
        flush_code()
        if not in_table:
            in_table = True
        table_rows.append(line.rstrip('\n'))
        i += 1
        continue
    else:
        if in_table:
            flush_table()
            in_table = False

    stripped = line.strip()

    # 标题
    if stripped.startswith('# ') and not stripped.startswith('## '):
        flush_code()
        flush_table()
        doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith('## '):
        flush_code()
        flush_table()
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('### '):
        flush_code()
        flush_table()
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith('---'):
        flush_code()
        flush_table()
        doc.add_paragraph('─' * 40)
    elif stripped.startswith('- ') or stripped.startswith('* '):
        flush_code()
        flush_table()
        text = stripped[2:]
        # 处理加粗
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        p = doc.add_paragraph(text, style='List Bullet')
    elif re.match(r'^\d+\.\s', stripped):
        flush_code()
        flush_table()
        text = re.sub(r'^\d+\.\s', '', stripped)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        p = doc.add_paragraph(text, style='List Number')
    elif stripped == '':
        flush_code()
        flush_table()
    else:
        flush_code()
        flush_table()
        # 普通段落，处理行内样式
        text = stripped
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # 链接
        if text:
            p = doc.add_paragraph(text)

    i += 1

# 清理
flush_code()
flush_table()

# 保存
output_path = 'g:/team_Software_Engineering_project/实验报告-Web UI接入与前后端分离.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
